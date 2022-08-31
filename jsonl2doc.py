import sys
import json
import docx
import tqdm

def yield_docs(fnames):
    for fname in fnames:
        with open(fname) as inp:
            for line in inp:
                doc=json.loads(line)
                doc["srcfile"]=fname
                ratio=len(doc["text"])/len(doc["summary"])
                if ratio<=10.0:
                    yield(doc)

if __name__=="__main__":
    
    doc=docx.Document()
    doc_counter=0
    current_len=0
    max_len=950000

    doclist=[]
    for d_idx, d in tqdm.tqdm(enumerate(yield_docs(sys.argv[1:]))):
        d_len=len(d["title"])+len(d["summary"])+len(d["text"])+50
        if (current_len+d_len)>max_len:
            #Save first!
            doc.save(f"xlsum-doc-in/xlsum-{doc_counter:05d}.docx")
            doc_counter+=1
            current_len=0
            doc=docx.Document()
        pgraph=doc.add_paragraph("")
        r=pgraph.add_run(f"Document number {d_idx}")
        doclist.append((d["srcfile"],d["id"]))
        r.bold=True
        r.underline=True
        pgraph=doc.add_paragraph(d["title"])
        pgraph=doc.add_paragraph(d["text"])
        
        pgraph=doc.add_paragraph("")
        r=pgraph.add_run(f"Summary")
        r.bold=True

        pgraph=doc.add_paragraph(d["summary"])
        current_len+=d_len
    else:
        doc.save(f"xlsum-doc-in/xlsum-{doc_counter:05d}.docx")
    with open("xlsum-doc-in/z_metadata.json","wt") as f:
        json.dump(doclist,f)
        
        
            
        
